"""Document upload, listing, and deletion endpoints."""

from fastapi import APIRouter, Depends, HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.ai.rag import delete_document_chunks, index_document
from app.config import settings
from app.database import get_db
from app.models import Document
from app.schemas import DocumentOut

router = APIRouter(prefix="/documents", tags=["documents"])

SUMMARY_SYSTEM = """\
You are an AI that summarizes documents concisely.
Write a 2-4 sentence summary capturing the main topics, purpose, and key information.
Be factual and specific. Do not use markdown formatting.
"""


@router.post("/upload", response_model=DocumentOut, status_code=201)
async def upload_document(file: UploadFile, db: Session = Depends(get_db)):
    filename = file.filename or "untitled"
    content_bytes = await file.read()

    # Reject files larger than 20 MB
    if len(content_bytes) > 20 * 1024 * 1024:
        raise HTTPException(413, "File too large. Maximum size is 20 MB.")

    # Validate file type
    allowed_ext = (".txt", ".docx", ".pdf", ".md", ".csv")
    if not any(filename.lower().endswith(ext) for ext in allowed_ext):
        raise HTTPException(415, f"Only {', '.join(allowed_ext)} files are supported.")

    # Extract text pages
    try:
        pages = _extract_pages(filename, content_bytes)
    except Exception as e:
        raise HTTPException(400, f"Could not read file: {e}")

    text = "\n\n".join(t for _, t in pages)
    if not text.strip():
        raise HTTPException(400, "File appears to be empty.")

    # Generate AI summary
    try:
        summary = await _generate_summary(text)
    except Exception as e:
        raise HTTPException(500, f"AI summary failed: {e}")

    ext_to_type = {".txt": "txt", ".docx": "docx", ".pdf": "pdf", ".md": "md", ".csv": "csv"}
    file_type = next((t for ext, t in ext_to_type.items() if filename.endswith(ext)), "txt")

    # Insert SQL record
    doc = Document(
        filename=filename,
        file_type=file_type,
        summary=summary,
        char_count=len(text),
        chunk_count=0,  # updated after indexing
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Index into ChromaDB — rollback SQL record on failure
    try:
        chunk_count = index_document(doc.id, pages)
        doc.chunk_count = chunk_count
        db.commit()
        db.refresh(doc)
    except Exception as e:
        db.delete(doc)
        db.commit()
        raise HTTPException(500, f"Failed to index document: {e}")

    return doc


@router.get("/", response_model=list[DocumentOut])
def list_documents(db: Session = Depends(get_db)):
    return db.query(Document).order_by(Document.created_at.desc()).all()


@router.delete("/{doc_id}", status_code=204)
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, f"Document #{doc_id} not found.")

    try:
        delete_document_chunks(doc_id)
    except Exception as e:
        raise HTTPException(500, f"Failed to remove document chunks: {e}")
    db.delete(doc)
    db.commit()


# ── Helpers ───────────────────────────────────────────────────────────────────


def _extract_pages(filename: str, content: bytes) -> list[tuple[int, str]]:
    """Extract text as (page_num, text) pairs.

    PDFs get real 1-indexed page numbers from pypdf.
    TXT and DOCX are treated as a single page (page 1).
    """
    if filename.endswith(".txt") or filename.endswith(".md"):
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")
        return [(1, text)]

    if filename.endswith(".csv"):
        import csv as csv_mod
        import io

        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")
        reader = csv_mod.reader(io.StringIO(text))
        rows = [
            "| " + " | ".join(cell.strip() for cell in row) + " |"
            for row in reader
            if any(c.strip() for c in row)
        ]
        if rows:
            # Insert separator after header row
            header = rows[0]
            col_count = header.count("|") - 1
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            rows.insert(1, separator)
        table_text = "<!-- TABLE -->\n" + "\n".join(rows) + "\n<!-- /TABLE -->"
        return [(1, table_text)]

    if filename.endswith(".docx"):
        import io

        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # Extract tables as markdown with markers
        table_blocks: list[str] = []
        for table in doc.tables:
            seen_rows: set[str] = set()
            md_rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                row_text = "| " + " | ".join(cells) + " |"
                key = " | ".join(cells)
                if key not in seen_rows and key.replace("|", "").strip():
                    seen_rows.add(key)
                    md_rows.append(row_text)
            if md_rows:
                # Insert separator after header row
                col_count = len(table.rows[0].cells) if table.rows else 1
                separator = "| " + " | ".join(["---"] * col_count) + " |"
                md_rows.insert(1, separator)
                table_blocks.append(
                    "<!-- TABLE -->\n" + "\n".join(md_rows) + "\n<!-- /TABLE -->"
                )
        if table_blocks:
            text = text + "\n\n" + "\n\n".join(table_blocks)

        return [(1, text)]

    if filename.endswith(".pdf"):
        import fitz  # PyMuPDF

        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for i, page in enumerate(doc, start=1):
            text = _extract_pdf_page(page)
            if not text:
                # OCR fallback for image-only pages (requires Tesseract)
                try:
                    tp = page.get_textpage_ocr(language="eng", dpi=300)
                    text = page.get_text(textpage=tp).strip()
                except Exception:
                    pass  # Tesseract not installed — skip this page
            if text:
                pages.append((i, text))
        doc.close()
        return pages

    raise ValueError(f"Unsupported file type: {filename}")


def _extract_pdf_page(page) -> str:  # noqa: ANN001
    """Layout-aware PDF page extraction with table and heading detection.

    Uses PyMuPDF's find_tables() for tables and get_text("dict") for
    heading detection. Falls back to plain get_text() on any error.
    """
    try:
        return _extract_pdf_page_inner(page)
    except Exception:
        return page.get_text().strip()


def _extract_pdf_page_inner(page) -> str:  # noqa: ANN001
    """Inner implementation — may raise; caller handles fallback."""
    import statistics

    # 1. Detect tables and format as markdown
    table_blocks: list[tuple[float, str]] = []  # (y_position, markdown)
    table_rects: list[tuple[float, float, float, float]] = []

    tables = page.find_tables()
    for table in tables:
        bbox = table.bbox  # (x0, y0, x1, y1)
        table_rects.append(bbox)
        rows = table.extract()
        if not rows:
            continue
        # Filter out fully-empty rows
        rows = [r for r in rows if any((c or "").strip() for c in r)]
        if not rows:
            continue
        # Build markdown table
        md_rows = []
        for row in rows:
            cells = [(c or "").strip().replace("\n", " ") for c in row]
            md_rows.append("| " + " | ".join(cells) + " |")
        # Insert separator after header
        col_count = len(rows[0])
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        md_rows.insert(1, separator)
        block_text = "<!-- TABLE -->\n" + "\n".join(md_rows) + "\n<!-- /TABLE -->"
        table_blocks.append((bbox[1], block_text))  # sort by y0

    # 2. Extract non-table text with heading detection
    def _overlaps_table(block_rect: tuple) -> bool:
        bx0, by0, bx1, by1 = block_rect[:4]
        for tx0, ty0, tx1, ty1 in table_rects:
            # Check for overlap (not just containment)
            if bx0 < tx1 and bx1 > tx0 and by0 < ty1 and by1 > ty0:
                return True
        return False

    page_dict = page.get_text("dict")
    blocks = page_dict.get("blocks", [])

    # Collect all font sizes to find median
    all_sizes: list[float] = []
    for block in blocks:
        if block.get("type") != 0:  # text blocks only
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if span.get("text", "").strip():
                    all_sizes.append(span["size"])

    median_size = statistics.median(all_sizes) if all_sizes else 12.0
    heading_threshold = median_size * 1.3

    text_blocks: list[tuple[float, str]] = []  # (y_position, text)
    for block in blocks:
        if block.get("type") != 0:
            continue
        bbox = block.get("bbox", (0, 0, 0, 0))
        if _overlaps_table(bbox):
            continue

        lines: list[str] = []
        for line in block.get("lines", []):
            spans_text = []
            is_heading = False
            for span in line.get("spans", []):
                text = span.get("text", "")
                if text.strip():
                    spans_text.append(text)
                    if span["size"] >= heading_threshold:
                        is_heading = True
            line_text = "".join(spans_text).strip()
            if line_text:
                if is_heading:
                    line_text = "<!-- HEADING -->" + line_text
                lines.append(line_text)

        if lines:
            text_blocks.append((bbox[1], "\n".join(lines)))

    # 3. Merge all blocks in reading order (by vertical position)
    all_blocks = table_blocks + text_blocks
    all_blocks.sort(key=lambda b: b[0])

    return "\n\n".join(text for _, text in all_blocks).strip()


async def _generate_summary(text: str) -> str:
    # Truncate to first 4000 chars to stay within token limits
    snippet = text[:4000]
    if settings.AI_PROVIDER == "anthropic":
        return await _summarize_anthropic(snippet)
    return await _summarize_openai(snippet)


async def _summarize_openai(text: str) -> str:
    from openai import AsyncOpenAI

    client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
    response = await client.chat.completions.create(
        model=settings.AI_MODEL,
        messages=[
            {"role": "system", "content": SUMMARY_SYSTEM},
            {"role": "user", "content": f"Summarize this document:\n\n{text}"},
        ],
        temperature=0.2,
        max_tokens=300,
    )
    return response.choices[0].message.content or "No summary available."


async def _summarize_anthropic(text: str) -> str:
    from anthropic import AsyncAnthropic

    client = AsyncAnthropic(api_key=settings.ANTHROPIC_API_KEY)
    response = await client.messages.create(
        model=settings.AI_MODEL,
        max_tokens=300,
        system=SUMMARY_SYSTEM,
        messages=[{"role": "user", "content": f"Summarize this document:\n\n{text}"}],
        temperature=0.2,
    )
    return response.content[0].text if response.content else "No summary available."
